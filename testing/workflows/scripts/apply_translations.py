#!/usr/bin/env python3
"""Apply translations from JSON back into a DOCX file."""
import json
import re
import sys
from pathlib import Path
from docx import Document

BODY_RE = re.compile(r"^body:p(\d+):r(\d+)$")
HEADER_RE = re.compile(r"^header:s(\d+):p(\d+):r(\d+)$")
FOOTER_RE = re.compile(r"^footer:s(\d+):p(\d+):r(\d+)$")

def apply_translations(docx_path: str, translations_path: str, output_path: str):
    doc = Document(docx_path)
    translations = json.loads(Path(translations_path).read_text(encoding="utf-8"))
    
    applied = 0
    for item in translations:
        unit_id = item["id"]
        target = item.get("target", item.get("source", ""))
        
        run = None
        m = BODY_RE.match(unit_id)
        if m:
            p_i, r_i = map(int, m.groups())
            try:
                run = doc.paragraphs[p_i].runs[r_i]
            except IndexError:
                continue
        
        m = HEADER_RE.match(unit_id)
        if m:
            s_i, p_i, r_i = map(int, m.groups())
            try:
                run = doc.sections[s_i].header.paragraphs[p_i].runs[r_i]
            except IndexError:
                continue
        
        m = FOOTER_RE.match(unit_id)
        if m:
            s_i, p_i, r_i = map(int, m.groups())
            try:
                run = doc.sections[s_i].footer.paragraphs[p_i].runs[r_i]
            except IndexError:
                continue
        
        if run:
            run.text = target
            applied += 1
    
    doc.save(output_path)
    print(f"Applied {applied} translations -> {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python apply_translations.py <input.docx> <translations.json> <output.docx>")
        sys.exit(1)
    
    apply_translations(sys.argv[1], sys.argv[2], sys.argv[3])
