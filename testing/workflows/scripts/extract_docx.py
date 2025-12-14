#!/usr/bin/env python3
"""Extract translatable text units from a DOCX file to JSON."""
import json
import re
import sys
from pathlib import Path
from docx import Document

CODE_PATTERN = re.compile(r"^\s*\d{3}-\d{5}[A-Z]?\s*$|^\s*\d{6}\s*$")

def extract_units(docx_path: str) -> list:
    doc = Document(docx_path)
    units = []
    
    # Body paragraphs
    for pi, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ""
        for ri, run in enumerate(p.runs):
            text = run.text
            if text and text.strip() and not CODE_PATTERN.match(text.strip()):
                units.append({
                    "id": f"body:p{pi}:r{ri}",
                    "source": text,
                    "style": style,
                    "where": "body"
                })
    
    # Headers/Footers
    for si, section in enumerate(doc.sections):
        for pi, p in enumerate(section.header.paragraphs):
            for ri, run in enumerate(p.runs):
                text = run.text
                if text and text.strip():
                    units.append({
                        "id": f"header:s{si}:p{pi}:r{ri}",
                        "source": text,
                        "style": p.style.name if p.style else "",
                        "where": "header"
                    })
        
        for pi, p in enumerate(section.footer.paragraphs):
            for ri, run in enumerate(p.runs):
                text = run.text
                if text and text.strip():
                    units.append({
                        "id": f"footer:s{si}:p{pi}:r{ri}",
                        "source": text,
                        "style": p.style.name if p.style else "",
                        "where": "footer"
                    })
    
    return units

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python extract_docx.py <input.docx> <output.json>")
        sys.exit(1)
    
    units = extract_units(sys.argv[1])
    Path(sys.argv[2]).write_text(json.dumps(units, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Extracted {len(units)} units -> {sys.argv[2]}")
